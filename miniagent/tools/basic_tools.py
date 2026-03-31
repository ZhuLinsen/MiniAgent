"""
Basic tools module providing common utility functions.
"""

import re
import requests
import math
import os
import platform
import subprocess
import datetime
import psutil
import webbrowser
from typing import Dict, List, Any, Optional
from pathlib import Path

from . import register_tool
from ..logger import get_logger

logger = get_logger(__name__)

@register_tool
def calculator(expression: str) -> Dict[str, Any]:
    """
    Calculate the result of a mathematical expression.
    Uses AST parsing for safety — only allows arithmetic and math functions.

    Args:
        expression: The mathematical expression to calculate, e.g., "2 + 2 * 3", "sqrt(16)", "sin(pi/2)"

    Returns:
        Dictionary containing the expression and result
    """
    import ast as _ast

    safe_dict = {
        'abs': abs, 'pow': pow, 'round': round, 'min': min, 'max': max,
        'int': int, 'float': float,
        'sin': math.sin, 'cos': math.cos, 'tan': math.tan,
        'asin': math.asin, 'acos': math.acos, 'atan': math.atan,
        'sinh': math.sinh, 'cosh': math.cosh, 'tanh': math.tanh,
        'exp': math.exp, 'log': math.log, 'log10': math.log10,
        'sqrt': math.sqrt, 'ceil': math.ceil, 'floor': math.floor,
        'degrees': math.degrees, 'radians': math.radians,
        'pi': math.pi, 'e': math.e,
    }
    allowed_funcs = set(safe_dict.keys()) - {'pi', 'e'}

    def _check_node(node):
        if isinstance(node, _ast.Call):
            if not (isinstance(node.func, _ast.Name) and node.func.id in allowed_funcs):
                raise ValueError(f"Function not allowed: {_ast.dump(node.func)}")
        for child in _ast.iter_child_nodes(node):
            _check_node(child)

    try:
        expression = expression.strip()
        tree = _ast.parse(expression, mode='eval')
        _check_node(tree.body)
        result = eval(compile(tree, '<string>', 'eval'), {"__builtins__": {}}, safe_dict)
        return {"expression": expression, "result": result}
    except Exception as e:
        return {"expression": expression, "error": f"Error: {str(e)}"}

@register_tool
def get_current_time() -> Dict[str, Any]:
    """
    Get current time information.
    
    Returns:
        Detailed information about the current time, including ISO format, year, month, day, etc.
    """
    logger.info("get_current_time called")
    now = datetime.datetime.now()
    return {
        "iso": now.isoformat(),
        "year": now.year,
        "month": now.month,
        "day": now.day,
        "hour": now.hour,
        "minute": now.minute,
        "second": now.second,
        "weekday": now.strftime("%A"),
        "formatted": now.strftime("%Y-%m-%d %H:%M:%S")
    }

@register_tool
def system_info() -> Dict[str, Any]:
    """
    Get detailed information about the system.
    
    Returns:
        Dictionary with system information like OS, version, architecture, etc.
    """
    logger.info("system_info called")
    try:
        info = {
            "platform": platform.system(),
            "platform_release": platform.release(),
            "platform_version": platform.version(),
            "architecture": platform.machine(),
            "processor": platform.processor(),
            "hostname": platform.node(),
            "python_version": platform.python_version(),
            "time": datetime.datetime.now().isoformat()
        }
        
        # Add more detailed information for specific platforms
        if platform.system() == "Linux":
            try:
                # Get distribution information
                import distro
                info["distribution"] = distro.name(pretty=True)
                info["distribution_version"] = distro.version()
                info["distribution_codename"] = distro.codename()
            except ImportError:
                # Fallback if distro module is not available
                try:
                    with open('/etc/os-release', 'r') as f:
                        os_release = dict(line.strip().split('=', 1) for line in f if '=' in line)
                    info["distribution"] = os_release.get('PRETTY_NAME', '').strip('"')
                except (OSError, ValueError):
                    pass
                    
        return info
    except Exception as e:
        logger.error(f"Error getting system info: {str(e)}")
        raise ValueError(f"Failed to get system information: {str(e)}")

@register_tool
def file_stats(directory: str = ".", pattern: str = "*") -> Dict[str, Any]:
    """
    Get statistics about files in a directory.
    
    Args:
        directory: Directory path to analyze, default is current directory
        pattern: File pattern to match, e.g., "*.py" for Python files
        
    Returns:
        Statistics about files including count, total size, extensions, etc.
    """
    logger.info(f"file_stats called: {directory}/{pattern}")
    try:
        # Resolve the path to handle relative paths and symlinks
        path = Path(directory).resolve()
        
        if not path.exists():
            raise ValueError(f"Directory '{directory}' does not exist")
            
        if not path.is_dir():
            raise ValueError(f"'{directory}' is not a directory")
            
        # Get all files matching the pattern
        files = list(path.glob(pattern))
        
        files = [f for f in files if f.is_file()]
        
        # Calculate statistics
        total_size = sum(f.stat().st_size for f in files)
        
        # Count files by extension
        extensions = {}
        for file in files:
            ext = file.suffix.lower()
            if ext in extensions:
                extensions[ext] += 1
            else:
                extensions[ext] = 1
                
        # Get oldest and newest files
        if files:
            oldest_file = min(files, key=lambda f: f.stat().st_mtime)
            newest_file = max(files, key=lambda f: f.stat().st_mtime)
            oldest = {
                "path": str(oldest_file.relative_to(path.parent)),
                "modified": datetime.datetime.fromtimestamp(oldest_file.stat().st_mtime).isoformat()
            }
            newest = {
                "path": str(newest_file.relative_to(path.parent)),
                "modified": datetime.datetime.fromtimestamp(newest_file.stat().st_mtime).isoformat()
            }
        else:
            oldest = newest = None
            
        return {
            "directory": str(path),
            "pattern": pattern,
            "file_count": len(files),
            "total_size_bytes": total_size,
            "total_size_human": _format_size(total_size),
            "extensions": extensions,
            "oldest_file": oldest,
            "newest_file": newest,
            "analyzed_at": datetime.datetime.now().isoformat()
        }
    except Exception as e:
        logger.error(f"Error analyzing files: {str(e)}")
        raise ValueError(f"Failed to analyze files in '{directory}': {str(e)}")

def _format_size(size_bytes: int) -> str:
    """Helper function to format size in human-readable format"""
    if size_bytes == 0:
        return "0 B"
    size_names = ("B", "KB", "MB", "GB", "TB", "PB")
    i = int(math.floor(math.log(size_bytes, 1024)))
    p = math.pow(1024, i)
    s = round(size_bytes / p, 2)
    return f"{s} {size_names[i]}"

@register_tool
def web_search(query: str, num_results: int = 5) -> List[Dict[str, str]]:
    """
    Perform a web search using DuckDuckGo API.
    
    Args:
        query: Search query content
        num_results: Number of results to return, default is 5
        
    Returns:
        List of search results, each containing title, link, and snippet
    """
    logger.info(f"web_search called: {query}")
    try:
        # DuckDuckGo search API endpoint
        url = "https://serpapi.com/search"
        
        # Get API key from environment
        api_key = os.environ.get("SERPAPI_KEY")
        if not api_key:
            return [{"error": "SERPAPI_KEY environment variable not set. Please set it in your .env file."}]
        
        # Parameters for the search
        params = {
            'engine': 'duckduckgo',
            'q': query,
            'api_key': api_key,
            'kl': 'us-en'  # Region and language
        }
        
        # Send request
        response = requests.get(url, params=params, timeout=30)
        response.raise_for_status()
        
        # Parse the response
        data = response.json()
        
        # Extract organic results
        results = []
        if 'organic_results' in data:
            for result in data['organic_results'][:num_results]:
                results.append({
                    "title": result.get('title', ''),
                    "link": result.get('link', ''),
                    "snippet": result.get('snippet', '')
                })
        
        # Add knowledge graph if available
        if 'knowledge_graph' in data and len(results) < num_results:
            kg = data['knowledge_graph']
            results.append({
                "title": kg.get('title', ''),
                "link": kg.get('website', ''),
                "snippet": kg.get('description', '')
            })
        
        # Add related searches if needed
        if 'related_searches' in data and len(results) < num_results:
            for related in data['related_searches'][:num_results - len(results)]:
                results.append({
                    "title": f"Related: {related.get('query', '')}",
                    "link": related.get('link', ''),
                    "snippet": "Related search suggestion"
                })
        
        return results[:num_results]
    except Exception as e:
        logger.error(f"Search error: {str(e)}")
        return [{"error": f"Failed to execute search '{query}': {str(e)}"}]

@register_tool
def http_request(
    url: str, 
    method: str = "GET", 
    headers: Optional[Dict[str, str]] = None, 
    data: Optional[Dict[str, Any]] = None
) -> Dict[str, Any]:
    """
    Send an HTTP request and return the response.
    
    Args:
        url: Request URL (must be a public URL, private/internal addresses are blocked)
        method: Request method, default is "GET"
        headers: Request headers, default is None
        data: Request data, default is None
        
    Returns:
        HTTP response data including status code, headers, and body
    """
    import ipaddress
    from urllib.parse import urlparse
    import socket

    logger.info(f"http_request called: {method} {url}")

    # --- SSRF protection: block requests to private/internal networks ---
    try:
        parsed = urlparse(url)
        hostname = parsed.hostname
        if not hostname:
            raise ValueError("Invalid URL: no hostname")
        # Resolve hostname to IP and check if it's private
        resolved_ips = socket.getaddrinfo(hostname, parsed.port or 80, proto=socket.IPPROTO_TCP)
        for _family, _type, _proto, _canonname, sockaddr in resolved_ips:
            ip = ipaddress.ip_address(sockaddr[0])
            if ip.is_private or ip.is_loopback or ip.is_link_local or ip.is_reserved:
                raise ValueError(
                    f"Blocked request to private/internal address: {hostname} -> {ip}"
                )
    except (socket.gaierror, ValueError) as e:
        raise ValueError(f"URL validation failed: {e}")

    try:
        # Prepare request parameters
        kwargs = {
            "headers": headers or {}
        }
        
        # Handle different request methods with data
        if method.upper() in ["POST", "PUT", "PATCH"] and data:
            kwargs["json"] = data
            
        # Send request with certificate verification
        response = requests.request(method.upper(), url, timeout=30, verify=True, **kwargs)
        
        # Prepare response data
        try:
            response_data = response.json()
        except (ValueError, TypeError):
            response_data = response.text
            
        return {
            "status_code": response.status_code,
            "headers": dict(response.headers),
            "data": response_data
        }
    except Exception as e:
        raise ValueError(f"Failed to execute HTTP request: {str(e)}")

@register_tool
def disk_usage(path: str = "/") -> Dict[str, Any]:
    """
    Get disk usage information for a specified path.
    
    Args:
        path: Path to check disk usage, default is root directory
        
    Returns:
        Dictionary containing disk usage information
    """
    logger.info(f"disk_usage called: {path}")
    try:
        usage = psutil.disk_usage(path)
        return {
            "path": path,
            "total_bytes": usage.total,
            "used_bytes": usage.used,
            "free_bytes": usage.free,
            "percent_used": usage.percent,
            "total_human": _format_size(usage.total),
            "used_human": _format_size(usage.used),
            "free_human": _format_size(usage.free),
            "updated_at": datetime.datetime.now().isoformat()
        }
    except Exception as e:
        logger.error(f"Error getting disk usage: {str(e)}")
        raise ValueError(f"Failed to get disk usage for '{path}': {str(e)}")

@register_tool
def process_list(limit: int = 10) -> List[Dict[str, Any]]:
    """
    Get list of running processes.
    
    Args:
        limit: Maximum number of processes to return, sorted by CPU usage
        
    Returns:
        List of process information dictionaries
    """
    logger.info(f"process_list called (limit: {limit})")
    try:
        processes = []
        for proc in psutil.process_iter(['pid', 'name', 'cpu_percent', 'memory_percent', 'username']):
            try:
                pinfo = proc.info
                processes.append({
                    "pid": pinfo['pid'],
                    "name": pinfo['name'],
                    "cpu_percent": pinfo['cpu_percent'],
                    "memory_percent": pinfo['memory_percent'],
                    "username": pinfo['username']
                })
            except (psutil.NoSuchProcess, psutil.AccessDenied):
                continue
        
        # Sort by CPU usage and limit results
        processes.sort(key=lambda x: x['cpu_percent'], reverse=True)
        return processes[:limit]
    except Exception as e:
        logger.error(f"Error getting process list: {str(e)}")
        raise ValueError(f"Failed to get process list: {str(e)}")

@register_tool
def system_load() -> Dict[str, Any]:
    """
    Get system load information including CPU, memory, and disk usage.
    
    Returns:
        Dictionary containing system load information
    """
    logger.info("system_load called")
    try:
        # CPU usage
        cpu_percent = psutil.cpu_percent(interval=1)
        cpu_count = psutil.cpu_count()
        
        # Memory usage
        memory = psutil.virtual_memory()
        
        # Disk usage
        disk = psutil.disk_usage('/')
        
        return {
            "cpu": {
                "percent": cpu_percent,
                "count": cpu_count,
                "load_avg": psutil.getloadavg() if hasattr(psutil, "getloadavg") else None
            },
            "memory": {
                "total": memory.total,
                "available": memory.available,
                "used": memory.used,
                "free": memory.free,
                "percent": memory.percent,
                "total_human": _format_size(memory.total),
                "available_human": _format_size(memory.available),
                "used_human": _format_size(memory.used),
                "free_human": _format_size(memory.free)
            },
            "disk": {
                "total": disk.total,
                "used": disk.used,
                "free": disk.free,
                "percent": disk.percent,
                "total_human": _format_size(disk.total),
                "used_human": _format_size(disk.used),
                "free_human": _format_size(disk.free)
            },
            "updated_at": datetime.datetime.now().isoformat()
        }
    except Exception as e:
        logger.error(f"Error getting system load: {str(e)}")
        raise ValueError(f"Failed to get system load information: {str(e)}")


@register_tool
def open_browser(url: str) -> str:
    """
    Open a URL in the default web browser.
    
    Args:
        url: The URL to open. Can also be a search query (will use Google).
        
    Returns:
        Status message
    """
    # If it doesn't look like a URL, treat it as a search query
    if not url.startswith(('http://', 'https://', 'file://')):
        from urllib.parse import quote
        url = f"https://www.google.com/search?q={quote(url)}"
    
    try:
        webbrowser.open(url)
        return f"Opened browser with: {url}"
    except Exception as e:
        raise ValueError(f"Failed to open browser: {str(e)}")


@register_tool
def open_app(app_name: str) -> str:
    """
    Open an application by name.
    
    Args:
        app_name: Name of the application to open (e.g., 'notepad', 'code', 'chrome')
        
    Returns:
        Status message
    """
    system = platform.system()
    
    try:
        if system == "Windows":
            # Common Windows app mappings
            app_map = {
                'notepad': 'notepad.exe',
                'calculator': 'calc.exe',
                'explorer': 'explorer.exe',
                'cmd': 'cmd.exe',
                'powershell': 'powershell.exe',
                'chrome': 'chrome',
                'firefox': 'firefox',
                'edge': 'msedge',
                'code': 'code',
                'vscode': 'code',
            }
            cmd = app_map.get(app_name.lower(), app_name)
            subprocess.Popen([cmd])
        elif system == "Darwin":  # macOS
            subprocess.Popen(['open', '-a', app_name])
        else:  # Linux
            subprocess.Popen([app_name], start_new_session=True)
        
        return f"Opened application: {app_name}"
    except Exception as e:
        raise ValueError(f"Failed to open application '{app_name}': {str(e)}")


@register_tool  
def clipboard_copy(text: str) -> str:
    """
    Copy text to system clipboard.
    
    Args:
        text: Text to copy to clipboard
        
    Returns:
        Status message
    """
    system = platform.system()
    
    try:
        if system == "Windows":
            subprocess.run(['clip'], input=text.encode('utf-16le'), check=True)
        elif system == "Darwin":  # macOS
            subprocess.run(['pbcopy'], input=text.encode('utf-8'), check=True)
        else:  # Linux
            # Try xclip first, then xsel
            try:
                subprocess.run(['xclip', '-selection', 'clipboard'], input=text.encode('utf-8'), check=True)
            except FileNotFoundError:
                subprocess.run(['xsel', '--clipboard', '--input'], input=text.encode('utf-8'), check=True)
        
        return f"Copied {len(text)} characters to clipboard"
    except Exception as e:
        raise ValueError(f"Failed to copy to clipboard: {str(e)}")


@register_tool
def clipboard_read() -> str:
    """
    Read text content from system clipboard.

    Returns:
        Text content from clipboard
    """
    system = platform.system()

    try:
        if system == "Windows":
            result = subprocess.run(['powershell', '-command', 'Get-Clipboard'], capture_output=True, text=True, check=True)
            return result.stdout.rstrip('\r\n')
        elif system == "Darwin":  # macOS
            result = subprocess.run(['pbpaste'], capture_output=True, text=True, check=True)
            return result.stdout
        else:  # Linux
            try:
                result = subprocess.run(['xclip', '-selection', 'clipboard', '-o'], capture_output=True, text=True, check=True)
                return result.stdout
            except FileNotFoundError:
                result = subprocess.run(['xsel', '--clipboard', '--output'], capture_output=True, text=True, check=True)
                return result.stdout
    except Exception as e:
        raise ValueError(f"Failed to read clipboard: {str(e)}")


@register_tool
def create_docx(path: str, content: str, title: str = "") -> str:
    """
    Create a Word document (.docx) with the given content.
    Note: Requires python-docx package. If not installed, creates a .txt file instead.
    
    Args:
        path: Path for the output file (will add .docx extension if needed)
        content: Text content for the document
        title: Optional title for the document
        
    Returns:
        Status message with file path
    """
    # Ensure .docx extension
    if not path.lower().endswith('.docx'):
        path = path + '.docx'
    
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        
        # Add title if provided
        if title:
            title_para = doc.add_heading(title, level=0)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Split content into paragraphs and add them
        paragraphs = content.split('\n\n')
        for para_text in paragraphs:
            if para_text.strip():
                # Check if it's a heading (starts with #)
                if para_text.startswith('# '):
                    doc.add_heading(para_text[2:], level=1)
                elif para_text.startswith('## '):
                    doc.add_heading(para_text[3:], level=2)
                elif para_text.startswith('### '):
                    doc.add_heading(para_text[4:], level=3)
                else:
                    doc.add_paragraph(para_text.strip())
        
        doc.save(path)
        return f"Created Word document: {path}"
        
    except ImportError:
        # Fallback: create a text file with instructions
        txt_path = path.replace('.docx', '.txt')
        with open(txt_path, 'w', encoding='utf-8') as f:
            if title:
                f.write(f"{title}\n{'='*len(title)}\n\n")
            f.write(content)
        return f"python-docx not installed. Created text file instead: {txt_path}\nTo create .docx files, run: pip install python-docx"


# Sensitive environment variable name patterns
_SENSITIVE_ENV_PATTERNS = re.compile(
    r"(KEY|SECRET|PASSWORD|TOKEN|CREDENTIAL|PRIVATE)", re.IGNORECASE
)


@register_tool
def env_get(name: str) -> str:
    """
    Get an environment variable value.
    Sensitive variables (containing KEY, SECRET, PASSWORD, TOKEN, etc.) are blocked.
    
    Args:
        name: Name of the environment variable
        
    Returns:
        Value of the environment variable or empty string if not found
    """
    if _SENSITIVE_ENV_PATTERNS.search(name):
        return "[blocked] Cannot read sensitive environment variables"
    return os.environ.get(name, "")


@register_tool
def env_set(name: str, value: str) -> str:
    """
    Set an environment variable (for current process only).
    Sensitive variables (containing KEY, SECRET, PASSWORD, TOKEN, etc.) are blocked.
    
    Args:
        name: Name of the environment variable
        value: Value to set
        
    Returns:
        Status message
    """
    if _SENSITIVE_ENV_PATTERNS.search(name):
        return "[blocked] Cannot modify sensitive environment variables"
    os.environ[name] = value
    return f"Set environment variable: {name}={value}"
